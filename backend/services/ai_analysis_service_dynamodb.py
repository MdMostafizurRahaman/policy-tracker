"""
DynamoDB-based AI Analysis service for policy analysis and scoring.
"""
import logging
import io
from datetime import datetime
from typing import Dict, Any, List, Optional
from fastapi import HTTPException
import uuid
import PyPDF2
from docx import Document
import openai
import os
from config.dynamodb import get_dynamodb
from config.data_constants import POLICY_AREAS
from utils.helpers import calculate_policy_score, calculate_completeness_score
from services.bedrock_service import bedrock_service

logger = logging.getLogger(__name__)

class AIAnalysisService:
    def __init__(self):
        # Service initialization
        pass
    
    def extract_text_from_file(self, file_content: bytes, filename: str) -> str:
        """Extract text content from uploaded file"""
        try:
            text_content = ""
            file_extension = filename.lower().split('.')[-1]
            
            logger.info(f"Extracting text from {filename} (type: {file_extension})")
            
            if file_extension == 'pdf':
                # Extract text from PDF
                try:
                    pdf_file = io.BytesIO(file_content)
                    pdf_reader = PyPDF2.PdfReader(pdf_file)
                    
                    logger.info(f"PDF has {len(pdf_reader.pages)} pages")
                    
                    for page_num, page in enumerate(pdf_reader.pages):
                        page_text = page.extract_text()
                        if page_text.strip():
                            text_content += page_text + "\n"
                            logger.info(f"Extracted {len(page_text)} characters from page {page_num + 1}")
                        else:
                            logger.warning(f"Page {page_num + 1} appears to be empty or image-based")
                    
                    if not text_content.strip():
                        logger.warning("PDF text extraction yielded no content - may be image-based PDF")
                        
                except Exception as pdf_error:
                    logger.error(f"PDF extraction error: {pdf_error}")
                    return ""
                    
            elif file_extension in ['doc', 'docx']:
                # Extract text from Word document
                try:
                    doc_file = io.BytesIO(file_content)
                    document = Document(doc_file)
                    
                    logger.info(f"Word document has {len(document.paragraphs)} paragraphs")
                    
                    for paragraph in document.paragraphs:
                        if paragraph.text.strip():
                            text_content += paragraph.text + "\n"
                    
                    # Also extract text from tables if any
                    for table in document.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                if cell.text.strip():
                                    text_content += cell.text + " "
                        text_content += "\n"
                        
                except Exception as doc_error:
                    logger.error(f"Word document extraction error: {doc_error}")
                    return ""
                    
            elif file_extension == 'txt':
                # Extract text from plain text file
                try:
                    # Try different encodings
                    encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
                    
                    for encoding in encodings:
                        try:
                            text_content = file_content.decode(encoding)
                            logger.info(f"Successfully decoded text file using {encoding}")
                            break
                        except UnicodeDecodeError:
                            continue
                    
                    if not text_content:
                        logger.error("Could not decode text file with any encoding")
                        return ""
                        
                except Exception as txt_error:
                    logger.error(f"Text file extraction error: {txt_error}")
                    return ""
                    
            else:
                logger.warning(f"Unsupported file type: {file_extension}")
                return ""
            
            # Clean and validate extracted text
            text_content = text_content.strip()
            
            if text_content:
                logger.info(f"Successfully extracted {len(text_content)} characters from {filename}")
                # Log first 200 characters for debugging
                logger.info(f"Text preview: {text_content[:200]}...")
            else:
                logger.warning(f"No text content extracted from {filename}")
                
            return text_content
            
        except Exception as e:
            logger.error(f"Text extraction failed for {filename}: {str(e)}")
            return ""
    
    def calculate_tea_scores(self, text_content: str) -> Dict[str, Any]:
        """
        Calculate Transparency, Explainability, Accountability (TEA) scores using Bedrock
        
        Args:
            text_content: The extracted text from the policy document
            
        Returns:
            Dictionary containing the three scores and detailed analysis
        """
        try:
            logger.info("Starting TEA scores calculation using Bedrock")
            
            if not text_content.strip():
                logger.warning("No text content provided for TEA analysis")
                return {
                    "transparency_score": 0,
                    "explainability_score": 0,
                    "accountability_score": 0,
                    "error": "No text content to analyze"
                }
            
            # Use Bedrock service to calculate scores
            scores_result = bedrock_service.calculate_transparency_explainability_accountability_scores(text_content)
            
            logger.info(f"TEA scores calculated: T={scores_result.get('scores', {}).get('transparency_score', 0)}, "
                       f"E={scores_result.get('scores', {}).get('explainability_score', 0)}, "
                       f"A={scores_result.get('scores', {}).get('accountability_score', 0)}")
            
            return scores_result
            
        except Exception as e:
            logger.error(f"TEA scores calculation failed: {str(e)}")
            return {
                "transparency_score": 0,
                "explainability_score": 0,
                "accountability_score": 0,
                "error": f"TEA analysis failed: {str(e)}"
            }
    
    def analyze_policy_document(self, text_content: str) -> Dict[str, Any]:
        """Analyze policy document content using GROQ API"""
        try:
            if not text_content.strip():
                logger.warning("No text content provided for analysis")
                return {"error": "No text content to analyze"}
            
            logger.info(f"Starting AI analysis of document with {len(text_content)} characters")
            
            # Use GROQ API instead of OpenAI
            groq_api_key = os.getenv('GROQ_API_KEY')
            groq_api_url = os.getenv('GROQ_API_URL')
            
            if not groq_api_key or not groq_api_url:
                logger.error("GROQ API credentials not configured")
                return {
                    "error": "GROQ API not configured",
                    "title": "Configuration Error",
                    "country": "Unknown",
                    "policy_area": "General",
                    "objectives": [],
                    "key_points": [],
                    "timeline": "Not specified",
                    "summary": "AI analysis not available - API not configured"
                }
            
            # Truncate text if too long and add sample text at the beginning for better analysis
            analysis_text = text_content[:3000] if len(text_content) > 3000 else text_content
            
            # Enhanced prompt for better extraction
            prompt = f"""
You are an expert policy analyst. Analyze the following policy document and extract specific information. 
Be precise and only extract information that is clearly present in the text.

DOCUMENT TEXT:
{analysis_text}

Please analyze and extract the following information in JSON format:

1. TITLE: The main title or name of the document/policy
2. COUNTRY: The country, jurisdiction, or region this policy applies to
3. POLICY_AREA: The main policy domain (e.g., "AI and Technology", "Healthcare", "Education", "Environment", etc.)
4. OBJECTIVES: List of main goals or objectives mentioned
5. KEY_POINTS: List of important policy points or provisions
6. TIMELINE: Any implementation dates or timelines mentioned
7. SUMMARY: A concise summary in 150 words or less

Return ONLY a valid JSON object with this exact structure:
{{
    "title": "exact title from document or descriptive title based on content",
    "country": "country name or Unknown if not found",
    "policy_area": "specific policy area",
    "objectives": ["objective 1", "objective 2"],
    "key_points": ["key point 1", "key point 2"],
    "timeline": "timeline information or Not specified",
    "summary": "comprehensive summary"
}}

Important: Return ONLY the JSON object, no additional text or explanations.
"""
            
            # Call GROQ API
            import requests
            headers = {
                'Authorization': f'Bearer {groq_api_key}',
                'Content-Type': 'application/json'
            }
            
            payload = {
                "model": "llama3-8b-8192",  # GROQ's model
                "messages": [
                    {"role": "system", "content": "You are a policy analysis expert. Extract key information from policy documents and return only valid JSON responses with no additional text."},
                    {"role": "user", "content": prompt}
                ],
                "max_tokens": 1500,  # Increased for better responses
                "temperature": 0.1,  # Lower temperature for more consistent extraction
                "top_p": 0.9
            }
            
            logger.info("Calling GROQ API for policy analysis")
            response = requests.post(groq_api_url, headers=headers, json=payload, timeout=30)
            
            if response.status_code != 200:
                logger.error(f"GROQ API error: {response.status_code} - {response.text}")
                raise Exception(f"GROQ API error: {response.status_code}")
            
            # Parse the response
            response_data = response.json()
            ai_response = response_data['choices'][0]['message']['content'].strip()
            
            logger.info(f"Received AI response: {ai_response[:200]}...")
            
            # Clean the response to extract JSON
            import json
            import re
            
            # Try to find JSON in the response
            json_match = re.search(r'\{.*\}', ai_response, re.DOTALL)
            if json_match:
                json_str = json_match.group()
            else:
                json_str = ai_response
            
            try:
                analysis_data = json.loads(json_str)
                logger.info("Successfully parsed AI analysis JSON")
                
                # Validate and ensure all required fields exist
                required_fields = ["title", "country", "policy_area", "objectives", "key_points", "timeline", "summary"]
                for field in required_fields:
                    if field not in analysis_data:
                        analysis_data[field] = "Not specified" if field in ["timeline", "summary"] else ("Unknown" if field in ["country", "policy_area", "title"] else [])
                
                # Extract meaningful data from the content if AI didn't find specific info
                if analysis_data.get("title") in ["Not specified", "Unknown", ""]:
                    # Try to extract title from first line or create one
                    lines = text_content.strip().split('\n')
                    first_line = lines[0].strip() if lines else ""
                    if first_line and len(first_line) < 100:
                        analysis_data["title"] = first_line
                    else:
                        analysis_data["title"] = "Policy Document"
                
                # If no country found, try to detect from common patterns
                if analysis_data.get("country") in ["Unknown", ""]:
                    country_patterns = [
                        r'(?i)\b(United States|USA|US|America)\b',
                        r'(?i)\b(United Kingdom|UK|Britain)\b',
                        r'(?i)\b(Canada|Canadian)\b',
                        r'(?i)\b(Australia|Australian)\b',
                        r'(?i)\b(Germany|German)\b',
                        r'(?i)\b(France|French)\b',
                        r'(?i)\b(Japan|Japanese)\b',
                        r'(?i)\b(China|Chinese)\b',
                        r'(?i)\b(India|Indian)\b',
                        r'(?i)\b(Bangladesh|BD)\b'
                    ]
                    
                    for pattern in country_patterns:
                        match = re.search(pattern, text_content)
                        if match:
                            country_map = {
                                'United States': 'United States', 'USA': 'United States', 'US': 'United States', 'America': 'United States',
                                'United Kingdom': 'United Kingdom', 'UK': 'United Kingdom', 'Britain': 'United Kingdom',
                                'Canada': 'Canada', 'Canadian': 'Canada',
                                'Australia': 'Australia', 'Australian': 'Australia',
                                'Germany': 'Germany', 'German': 'Germany',
                                'France': 'France', 'French': 'France',
                                'Japan': 'Japan', 'Japanese': 'Japan',
                                'China': 'China', 'Chinese': 'China',
                                'India': 'India', 'Indian': 'India',
                                'Bangladesh': 'Bangladesh', 'BD': 'Bangladesh'
                            }
                            analysis_data["country"] = country_map.get(match.group(), match.group())
                            break
                
                # Calculate TEA scores using Bedrock
                logger.info("Calculating TEA scores for the document")
                tea_scores = self.calculate_tea_scores(text_content)
                
                # Add TEA scores to analysis data
                analysis_data["tea_scores"] = tea_scores.get("scores", {
                    "transparency_score": 0,
                    "explainability_score": 0,
                    "accountability_score": 0
                })
                analysis_data["tea_analysis"] = {
                    "transparency_analysis": tea_scores.get("transparency_analysis", []),
                    "explainability_analysis": tea_scores.get("explainability_analysis", []),
                    "accountability_analysis": tea_scores.get("accountability_analysis", [])
                }
                
                # Add metadata about the analysis method
                analysis_data["tea_metadata"] = {
                    "analysis_method": "aws_bedrock" if not tea_scores.get("fallback_used") else "keyword_fallback",
                    "fallback_used": tea_scores.get("fallback_used", False),
                    "has_error": "error" in tea_scores
                }
                
                # Log the calculated scores
                scores = analysis_data["tea_scores"]
                logger.info(f"TEA Scores - Transparency: {scores.get('transparency_score', 0)}, "
                           f"Explainability: {scores.get('explainability_score', 0)}, "
                           f"Accountability: {scores.get('accountability_score', 0)}")
                
                if tea_scores.get("fallback_used"):
                    logger.info("TEA scores calculated using keyword-based fallback (Bedrock not available)")
                elif "error" in tea_scores:
                    logger.warning(f"TEA scores calculation had errors: {tea_scores.get('error')}")
                else:
                    logger.info("TEA scores calculated successfully using AWS Bedrock")
                
                return analysis_data
                
            except json.JSONDecodeError as e:
                logger.error(f"JSON parsing failed: {e}")
                logger.error(f"AI response was: {ai_response}")
                
                # Create structured response from raw text
                lines = text_content.strip().split('\n')
                title = lines[0].strip() if lines else "Policy Document"
                
                analysis_data = {
                    "title": title[:100],
                    "country": "Unknown", 
                    "policy_area": "General Policy",
                    "objectives": [],
                    "key_points": [],
                    "timeline": "Not specified",
                    "summary": f"Analysis of policy document. {ai_response[:150]}..."
                }
                
                return analysis_data
            
        except Exception as e:
            logger.error(f"Policy document analysis failed: {str(e)}")
            return {
                "error": f"Analysis failed: {str(e)}",
                "title": "Analysis Failed",
                "country": "Unknown",
                "policy_area": "General",
                "objectives": [],
                "key_points": [],
                "timeline": "Not specified", 
                "summary": "AI analysis could not be completed due to technical error"
            }
    
    async def _get_db(self):
        """Get DynamoDB client"""
        return await get_dynamodb()
    
    async def analyze_policy(self, policy_data: Dict[str, Any]) -> Dict[str, Any]:
        """Analyze a policy submission and provide AI insights"""
        try:
            logger.info("Starting AI policy analysis")
            
            # Calculate basic scores
            policy_score = calculate_policy_score(policy_data)
            completeness_score = calculate_completeness_score(policy_data)
            
            # Analyze policy areas coverage
            covered_areas = set()
            total_policies = 0
            
            for area in policy_data.get('policy_areas', []):
                covered_areas.add(area.get('area_name', ''))
                total_policies += len(area.get('policies', []))
            
            # Calculate coverage percentage
            total_possible_areas = len(POLICY_AREAS)
            coverage_percentage = (len(covered_areas) / total_possible_areas) * 100 if total_possible_areas > 0 else 0
            
            # Generate insights
            insights = []
            
            # Coverage insights
            if coverage_percentage >= 80:
                insights.append({
                    "type": "strength",
                    "category": "coverage",
                    "message": f"Excellent policy coverage across {len(covered_areas)} areas ({coverage_percentage:.1f}%)"
                })
            elif coverage_percentage >= 50:
                insights.append({
                    "type": "improvement",
                    "category": "coverage",
                    "message": f"Good policy coverage, but consider expanding to more areas. Currently covering {coverage_percentage:.1f}%"
                })
            else:
                insights.append({
                    "type": "weakness",
                    "category": "coverage",
                    "message": f"Limited policy coverage ({coverage_percentage:.1f}%). Consider adding policies in more areas"
                })
            
            # Policy quantity insights
            if total_policies >= 20:
                insights.append({
                    "type": "strength",
                    "category": "quantity",
                    "message": f"Strong policy portfolio with {total_policies} total policies"
                })
            elif total_policies >= 10:
                insights.append({
                    "type": "neutral",
                    "category": "quantity",
                    "message": f"Moderate policy portfolio with {total_policies} policies"
                })
            else:
                insights.append({
                    "type": "improvement",
                    "category": "quantity",
                    "message": f"Consider expanding policy portfolio. Currently {total_policies} policies"
                })
            
            # Missing areas analysis
            all_area_names = {area['name'] for area in POLICY_AREAS}
            missing_areas = all_area_names - covered_areas
            
            if missing_areas:
                top_missing = list(missing_areas)[:3]  # Show top 3 missing areas
                insights.append({
                    "type": "suggestion",
                    "category": "missing_areas",
                    "message": f"Consider adding policies in: {', '.join(top_missing)}"
                })
            
            # Quality insights based on completeness
            if completeness_score >= 80:
                insights.append({
                    "type": "strength",
                    "category": "quality",
                    "message": "High quality policy descriptions with good detail"
                })
            elif completeness_score >= 60:
                insights.append({
                    "type": "improvement",
                    "category": "quality",
                    "message": "Policy descriptions could be more detailed for better impact"
                })
            else:
                insights.append({
                    "type": "weakness",
                    "category": "quality",
                    "message": "Policy descriptions need significant improvement for clarity and impact"
                })
            
            # Generate recommendations
            recommendations = []
            
            if coverage_percentage < 70:
                recommendations.append("Expand policy coverage to include more diverse areas")
            
            if total_policies < 15:
                recommendations.append("Increase the number of policies for comprehensive governance")
            
            if completeness_score < 70:
                recommendations.append("Improve policy descriptions with more specific details and implementation plans")
            
            # Add area-specific recommendations
            priority_areas = ["Healthcare", "Education", "Economic Development", "Environment"]
            for area in priority_areas:
                if area not in covered_areas:
                    recommendations.append(f"Consider adding policies in {area} for comprehensive coverage")
            
            analysis_result = {
                "analysis_id": str(uuid.uuid4()),
                "policy_score": policy_score,
                "completeness_score": completeness_score,
                "coverage_percentage": round(coverage_percentage, 1),
                "total_policies": total_policies,
                "covered_areas": list(covered_areas),
                "missing_areas": list(missing_areas),
                "insights": insights,
                "recommendations": recommendations,
                "overall_rating": self._calculate_overall_rating(policy_score, completeness_score, coverage_percentage),
                "analyzed_at": datetime.utcnow().isoformat()
            }
            
            logger.info("AI policy analysis completed successfully")
            return analysis_result
            
        except Exception as e:
            logger.error(f"AI analysis error: {str(e)}")
            raise HTTPException(status_code=500, detail=f"Analysis failed: {str(e)}")
    
    def _calculate_overall_rating(self, policy_score: float, completeness_score: float, coverage_percentage: float) -> str:
        """Calculate overall policy rating"""
        combined_score = (policy_score + completeness_score + coverage_percentage) / 3
        
        if combined_score >= 85:
            return "Excellent"
        elif combined_score >= 70:
            return "Good"
        elif combined_score >= 55:
            return "Fair"
        elif combined_score >= 40:
            return "Needs Improvement"
        else:
            return "Poor"
    
    async def get_country_analysis(self, country: str) -> Dict[str, Any]:
        """Get comprehensive analysis for a country's policies"""
        try:
            db = await self._get_db()
            
            # Get all approved policies for the country
            all_policies = await db.scan_table('policies')
            country_policies = [
                p for p in all_policies 
                if p.get('country', '').lower() == country.lower() 
                and p.get('status') == 'approved'
            ]
            
            if not country_policies:
                return {
                    "country": country,
                    "total_policies": 0,
                    "message": "No approved policies found for this country",
                    "analyzed_at": datetime.utcnow().isoformat()
                }
            
            # Aggregate analysis
            total_policies = sum(p.get('total_policies', 0) for p in country_policies)
            avg_policy_score = sum(p.get('score', 0) for p in country_policies) / len(country_policies)
            avg_completeness = sum(p.get('completeness_score', 0) for p in country_policies) / len(country_policies)
            
            # Collect all covered areas
            all_covered_areas = set()
            area_policy_counts = {}
            
            for policy in country_policies:
                for area in policy.get('policy_areas', []):
                    area_name = area.get('area_name', '')
                    all_covered_areas.add(area_name)
                    area_policy_counts[area_name] = area_policy_counts.get(area_name, 0) + len(area.get('policies', []))
            
            # Calculate coverage
            total_possible_areas = len(POLICY_AREAS)
            coverage_percentage = (len(all_covered_areas) / total_possible_areas) * 100 if total_possible_areas > 0 else 0
            
            return {
                "country": country,
                "total_submissions": len(country_policies),
                "total_policies": total_policies,
                "average_policy_score": round(avg_policy_score, 2),
                "average_completeness_score": round(avg_completeness, 2),
                "coverage_percentage": round(coverage_percentage, 1),
                "covered_areas": list(all_covered_areas),
                "policies_by_area": area_policy_counts,
                "overall_rating": self._calculate_overall_rating(avg_policy_score, avg_completeness, coverage_percentage),
                "analyzed_at": datetime.utcnow().isoformat()
            }
            
        except Exception as e:
            logger.error(f"Country analysis error: {str(e)}")
            raise HTTPException(status_code=500, detail=f"Country analysis failed: {str(e)}")
    
    async def compare_countries(self, countries: List[str]) -> Dict[str, Any]:
        """Compare policy performance across multiple countries"""
        try:
            if len(countries) < 2:
                raise HTTPException(status_code=400, detail="At least 2 countries required for comparison")
            
            country_analyses = []
            for country in countries:
                analysis = await self.get_country_analysis(country)
                if analysis.get('total_policies', 0) > 0:
                    country_analyses.append(analysis)
            
            if len(country_analyses) < 2:
                return {
                    "message": "Insufficient data for comparison",
                    "countries_with_data": [a['country'] for a in country_analyses],
                    "compared_at": datetime.utcnow().isoformat()
                }
            
            # Sort by overall performance
            country_analyses.sort(key=lambda x: x.get('average_policy_score', 0), reverse=True)
            
            # Find best and worst performers
            best_performer = country_analyses[0]
            worst_performer = country_analyses[-1]
            
            # Calculate averages across all countries
            avg_policies = sum(c.get('total_policies', 0) for c in country_analyses) / len(country_analyses)
            avg_score = sum(c.get('average_policy_score', 0) for c in country_analyses) / len(country_analyses)
            avg_coverage = sum(c.get('coverage_percentage', 0) for c in country_analyses) / len(country_analyses)
            
            return {
                "countries_compared": len(country_analyses),
                "comparison_data": country_analyses,
                "best_performer": {
                    "country": best_performer['country'],
                    "score": best_performer.get('average_policy_score', 0)
                },
                "worst_performer": {
                    "country": worst_performer['country'],
                    "score": worst_performer.get('average_policy_score', 0)
                },
                "averages": {
                    "policies": round(avg_policies, 1),
                    "score": round(avg_score, 2),
                    "coverage": round(avg_coverage, 1)
                },
                "compared_at": datetime.utcnow().isoformat()
            }
            
        except HTTPException:
            raise
        except Exception as e:
            logger.error(f"Country comparison error: {str(e)}")
            raise HTTPException(status_code=500, detail=f"Country comparison failed: {str(e)}")

# Global instance
ai_analysis_service = AIAnalysisService()
